
# y0__xDcirSgqveAAhjB3RMg69yB-ROiJgcc6EUOSHCTjlfOBUBzXPo3Kw

import os
import base64
import requests
from pdf2image import convert_from_path
from docx import Document
import time


# === НАСТРОЙ ЭТИ ПАРАМЕТРЫ ===
IAM_TOKEN = 't1.9euelZqOlpqUzsvLncqPnZiUmc2Nze3rnpWak56XzsaWzZGJnpSMxorMzJPl8_cEO2E7-e8NLSh1_d3z90RpXjv57w0tKHX9zef1656Vmoqdx8vMns2OnYyJmpfKy4vO7_zF656Vmoqdx8vMns2OnYyJmpfKy4vO.ohCeytxtRHRJHNDl63No-kqf2bt-SrWuLuB6q6SZaK4OTIiMNtnSqJVKagsBRmJ0MiheUNs3ev50Vk8Y7ZPNCQ'
FOLDER_ID = 'b1gejpaoh25hcp76j3f5'
PDF_PATH = r'C:\Users\ivan_\PycharmProjects\PythonProject2\Протокол совещания руковод. ПЧУ г. Санкт-Петербург-разб.pdf'

TMP_IMG_DIR = 'ocr_pages'
OUTPUT_DOCX = 'ocr_result.docx'

# === Подготовка ===
os.makedirs(TMP_IMG_DIR, exist_ok=True)

# === Преобразуем PDF → PNG по страницам ===
pages = convert_from_path(PDF_PATH, dpi=200)
image_paths = []
for i, page in enumerate(pages):
    img_path = os.path.join(TMP_IMG_DIR, f'page_{i+1}.png')
    page.save(img_path, 'PNG')
    image_paths.append(img_path)

# === Функция распознавания текста с PNG через Vision ===
def ocr_image(path, max_retries=5):
    with open(path, 'rb') as f:
        img_base64 = base64.b64encode(f.read()).decode('utf-8')

    url = 'https://vision.api.cloud.yandex.net/vision/v1/batchAnalyze'
    headers = {
        'Authorization': f'Bearer {IAM_TOKEN}',
        'Content-Type': 'application/json'
    }
    payload = {
        "folderId": FOLDER_ID,
        "analyze_specs": [{
            "content": img_base64,
            "features": [{
                "type": "TEXT_DETECTION",
                "text_detection_config": {
                    "language_codes": ["ru"]
                }
            }]
        }]
    }

    delay = 101
    for attempt in range(max_retries):
        resp = requests.post(url, headers=headers, json=payload)

        if resp.status_code == 429:
            print(f"⚠️ API вернул 429 (Too Many Requests), попытка {attempt+1}/{max_retries}. Жду {delay} секунд...")
            time.sleep(delay)
            delay *= 2  # экспоненциальная задержка
            continue

        if resp.status_code != 200:
            print(f"❌ Ошибка API: {resp.status_code}")
            return f"[ошибка {resp.status_code}]"

        time.sleep(2)  # базовая пауза
        break

    # === Извлечение текста ===
    result = resp.json()
    text = ''
    try:
        pages = result['results'][0]['results'][0]['textDetection']['pages']
        for page in pages:
            for block in page.get('blocks', []):
                for line in block.get('lines', []):
                    words = [word.get('text', '') for word in line.get('words', [])]
                    text += ' '.join(words) + '\n'
    except Exception as e:
        print(f"[!] Ошибка извлечения: {e}")
        return "[ошибка извлечения текста]"
    return text.strip()

# === Основной цикл по страницам ===
doc = Document()
for i, path in enumerate(image_paths):
    print(f"🔍 Распознаю страницу {i+1}/{len(image_paths)}...")
    text = ocr_image(path)
    doc.add_heading(f'Страница {i+1}', level=2)
    if text:
        for line in text.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("[нет текста]")

# === Сохраняем результат ===
doc.save(OUTPUT_DOCX)
print(f"\n✅ Распознанный текст сохранён в: {OUTPUT_DOCX}")

# === Очистка временных файлов ===
for p in image_paths:
    os.remove(p)
os.rmdir(TMP_IMG_DIR)
