import os
import base64
import requests
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import time

# === НАСТРОЙКИ ===
IAM_TOKEN = 't1.9euelZrNl8iZlZmPmYyQjJ2KlJ2dzO3rnpWak56XzsaWzZGJnpSMxorMzJPl8_dCIV47-e9HRyFa_d3z9wJQWzv570dHIVr9zef1656Vmo2dncyPx53OyM6Tm5KNiZyK7_zF656Vmo2dncyPx53OyM6Tm5KNiZyK.dAp69zivkS2doHBa5Rl_XwgTjog9qOVjgsZB0jZOnjZl_4IJc9u8UrbGYyC3UYQOlJITL_BAjp7F6_eMaOBSCQ'
FOLDER_ID = 'b1gejpaoh25hcp76j3f5'
INPUT_FILE = r'C:\Users\ivan_\PycharmProjects\PythonProject2\data\photo_2025-07-29_09-25-52.jpg'

# Папка для результатов
RESULT_DIR = 'result'
TMP_DIR = os.path.join(RESULT_DIR, 'tmp')
os.makedirs(TMP_DIR, exist_ok=True)

# Определяем имя без расширения
basename = os.path.splitext(os.path.basename(INPUT_FILE))[0]
OUTPUT_DOCX = os.path.join(RESULT_DIR, f"{basename}.docx")

# === Функция OCR одной картинки ===
def ocr_image(path, max_retries=5):
    with open(path, 'rb') as f:
        img_base64 = base64.b64encode(f.read()).decode('utf-8')

    url = 'https://vision.api.cloud.yandex.net/vision/v1/batchAnalyze'
    headers = {'Authorization': f'Bearer {IAM_TOKEN}', 'Content-Type': 'application/json'}
    payload = {
        "folderId": FOLDER_ID,
        "analyze_specs": [{
            "content": img_base64,
            "features": [{"type": "DOCUMENT_TEXT_DETECTION", "text_detection_config": {"language_codes": ["de"]}}]
        }]
    }

    delay = 5
    for attempt in range(max_retries):
        resp = requests.post(url, headers=headers, json=payload)
        if resp.status_code == 429:
            print(f"⚠️ 429 Too Many Requests, жду {delay} сек...")
            time.sleep(delay)
            delay *= 2
            continue
        if resp.status_code != 200:
            return f"[ошибка {resp.status_code}]"
        break

    # Извлечение текста
    try:
        result = resp.json()
        pages = result['results'][0]['results'][0]['textDetection']['pages']
        text = ''
        for page in pages:
            for block in page.get('blocks', []):
                for line in block.get('lines', []):
                    words = [word.get('text', '') for word in line.get('words', [])]
                    text += ' '.join(words) + '\n'
        return text.strip()
    except Exception as e:
        print(f"[!] Ошибка извлечения: {e}")
        return "[ошибка извлечения текста]"

# === Основная логика ===
doc = Document()
file_ext = os.path.splitext(INPUT_FILE)[1].lower()
image_paths = []

if file_ext == '.pdf':
    pages = convert_from_path(INPUT_FILE, dpi=300)
    for i, page in enumerate(pages):
        img_path = os.path.join(TMP_DIR, f'page_{i+1}.png')
        page.save(img_path, 'PNG')
        image_paths.append(img_path)
else:
    image_paths.append(INPUT_FILE)  # если сразу PNG/JPG

for i, path in enumerate(image_paths):
    print(f"🔍 Распознаю страницу {i+1}/{len(image_paths)}...")
    text = ocr_image(path)
    doc.add_heading(f'Страница {i+1}', level=2)
    doc.add_paragraph(text if text else "[нет текста]")

# Сохранение результата
os.makedirs(RESULT_DIR, exist_ok=True)
doc.save(OUTPUT_DOCX)
print(f"\n✅ Результат сохранён: {OUTPUT_DOCX}")

# Очистка временных файлов
for p in image_paths:
    if TMP_DIR in p:
        os.remove(p)
