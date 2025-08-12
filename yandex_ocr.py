# python
"""Utility for performing OCR using Yandex Cloud Vision API.

The script accepts an input image, PDF or a directory containing such
files, sends each page to the Yandex Cloud Vision API and saves the
recognised text into separate DOCX files. On request the OCR output of
all processed files can be merged into single DOCX, TXT and CSV
documents.

Usage:
    python yandex_ocr.py [INPUT_PATH] [--output-dir DIR]

If the IAM token is missing it is automatically requested from Yandex
Cloud using the built-in OAuth token. A default folder ID is embedded in
the script and may be overridden via command line or environment
variable.
"""

from __future__ import annotations

import argparse
import base64
import os
import io
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import csv

import requests
from docx import Document
from pdf2image import convert_from_path
from PIL import Image, ImageEnhance
import logging
import sys

import api

# Secrets должны приходить из переменных окружения или CLI‑параметров.
# Плейсхолдеры не препятствуют запуску, но потребуют явного ввода значений.
DEFAULT_FOLDER_ID = os.getenv("YANDEX_FOLDER_ID", "b1gejpaoh25hcp76j3f5")


def preprocess_image(path: Path, tmp_dir: Path) -> Path:
    """Improve image quality for OCR and save into *tmp_dir*.

    The image is upscaled, converted to grayscale and sharpened. The
    processed image is saved as JPEG in the temporary directory and its
    path is returned.
    """
    logging.info("Preprocessing %s", path)
    img = Image.open(path)

    if max(img.size) < 1000:
        new_size = (img.width * 2, img.height * 2)
        img = img.resize(new_size, Image.LANCZOS)

    # convert to 8‑bit grayscale
    img = img.convert("L")

    # increase contrast and sharpness
    img = ImageEnhance.Contrast(img).enhance(1.8)
    img = ImageEnhance.Sharpness(img).enhance(2.0)

    # light binarisation (keep as grayscale)
    img = img.point(lambda x: 0 if x < 140 else 255, "L")

    tmp_dir.mkdir(parents=True, exist_ok=True)
    tmp_path = tmp_dir / f"{path.stem}.jpg"
    img.save(tmp_path, "JPEG", quality=90)
    logging.debug("Preprocessed image stored at %s", tmp_path)
    return tmp_path


def _extract_text_from_pages(pages: List[Dict[str, Any]]) -> str:
    """Helper to build text from Vision pages structure."""
    text_lines: List[str] = []
    for page in pages or []:
        for block in page.get("blocks", []):
            for line in block.get("lines", []):
                words = [w.get("text", "") for w in line.get("words", []) if w.get("text")]
                if words:
                    text_lines.append(" ".join(words))
    return "\n".join(text_lines).strip()


def _build_auth_headers(iam_token: Optional[str], api_key: Optional[str]) -> Tuple[Dict[str, str], str]:
    """Return headers and auth mode string."""
    if iam_token:
        return {"Authorization": f"Bearer {iam_token}", "Content-Type": "application/json"}, "iam"
    if api_key:
        return {"Authorization": f"Api-Key {api_key}", "Content-Type": "application/json"}, "api_key"
    return {}, "none"


def ocr_image(path: Path, iam_token: Optional[str], folder_id: str, api_key: Optional[str] = None) -> str:
    """Send an image to the Yandex Vision API and return extracted text."""
    logging.info("Requesting OCR for %s", path)

    headers, mode = _build_auth_headers(iam_token, api_key)
    if mode == "none":
        logging.error("No credentials provided. Use --iam-token or --api-key (or set env vars).")
        return "[auth error: no credentials]"

    if not folder_id:
        logging.error("Folder ID is not set. Provide --folder-id or set YANDEX_FOLDER_ID.")
        return "[config error: no folder id]"

    if path.stat().st_size > 1_000_000:
        logging.debug("Reducing size of %s before upload", path)
        img = Image.open(path)
        quality = 85
        while True:
            buf = io.BytesIO()
            img.save(buf, "JPEG", quality=quality)
            if buf.tell() <= 1_000_000 or quality <= 30:
                break
            quality -= 5
        with open(path, "wb") as fh:
            fh.write(buf.getvalue())

    with open(path, "rb") as fh:
        img_base64 = base64.b64encode(fh.read()).decode("utf-8")

    url = "https://vision.api.cloud.yandex.net/vision/v1/batchAnalyze"
    payload = {
        "folderId": folder_id,
        "analyze_specs": [
            {
                "content": img_base64,
                "features": [
                    {
                        "type": "TEXT_DETECTION",
                        "text_detection_config": {
                            "language_codes": ["*"]
                        },
                    }
                ],
            }
        ],
    }

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
    except Exception:
        logging.exception("HTTP request failed for %s", path)
        return "[request error]"
    if resp.status_code != 200:
        logging.error(
            "Yandex API returned %s for %s: %s",
            resp.status_code,
            path,
            resp.text.strip(),
        )
        return f"[error {resp.status_code}]"

    try:
        result = resp.json()
    except Exception:
        logging.exception("Response is not JSON for %s", path)
        return "[response parse error]"

    try:
        # Robust traversal of results
        all_text_parts: List[str] = []
        root_results = result.get("results", [])
        if not root_results:
            logging.warning("No 'results' in API response for %s: %s", path, result)
            return "[no results]"

        for spec_idx, spec in enumerate(root_results):
            inner_results = spec.get("results", [])
            if not inner_results:
                logging.warning("Empty 'results' for spec %d in %s", spec_idx, path)
                continue

            for item_idx, item in enumerate(inner_results):
                if "error" in item:
                    err = item.get("error", {})
                    code = err.get("code")
                    msg = err.get("message")
                    logging.error("API error for %s (spec %d item %d): %s %s", path, spec_idx, item_idx, code, msg)
                    continue

                pages = None
                if "textDetection" in item:
                    pages = item["textDetection"].get("pages", [])
                elif "textAnnotation" in item:
                    pages = item["textAnnotation"].get("pages", [])

                if pages:
                    text_part = _extract_text_from_pages(pages)
                    if text_part:
                        all_text_parts.append(text_part)
                else:
                    logging.warning(
                        "No 'pages' in item for %s (spec %d item %d). Keys: %s",
                        path, spec_idx, item_idx, list(item.keys())
                    )

        final_text = "\n".join(p for p in all_text_parts if p).strip()
        if not final_text:
            logging.warning("No text extracted for %s. Response summary keys: %s", path, list(result.keys()))
            return "[no text]"
        return final_text

    except Exception:
        logging.exception("Failed to extract text for %s", path)
        try:
            logging.debug("Raw response for %s: %s", path, resp.text[:2000])
        except Exception:
            pass
        return "[text extraction error]"


def resolve_input(path: Path) -> Path:
    """Return an existing *Path* for ``path``.

    If *path* does not exist and refers only to a file or directory name
    (without any parent components), the current working directory is
    searched recursively for a matching entry. This avoids an expensive
    recursive search when a full path is provided but does not exist.
    """

    if path.exists():
        return path

    if len(path.parts) == 1:
        matches = list(Path.cwd().rglob(path.name))
        if matches:
            return matches[0]

    raise FileNotFoundError(f"Input path '{path}' not found")


def find_input_files(path: Path) -> List[Path]:
    """Return a list of image/PDF files found under *path*."""
    path = resolve_input(path)
    if path.is_file():
        logging.info("Input is a single file: %s", path)
        return [path]

    logging.info("Searching for files under %s", path)
    exts = {".pdf", ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}
    files = [p for p in path.rglob("*") if p.suffix.lower() in exts]
    logging.info("Found %d files", len(files))
    return files


def process_file(
    input_file: Path,
    output_docx: Path,
    tmp_dir: Path,
    iam_token: Optional[str],
    folder_id: str,
    api_key: Optional[str],
) -> List[str]:
    """Process *input_file* and save the recognised text to *output_docx*.

    Returns a list of extracted text pieces for each processed page.
    """
    logging.info("Processing file %s", input_file)
    images: List[Path] = []

    if input_file.suffix.lower() == ".pdf":
        logging.info("Splitting PDF %s", input_file)
        for i, page in enumerate(convert_from_path(str(input_file), dpi=300), start=1):
            img_path = tmp_dir / f"page_{i}.png"
            page.save(img_path, "PNG")
            images.append(preprocess_image(img_path, tmp_dir))
    else:
        images.append(preprocess_image(input_file, tmp_dir))

    document = Document()
    page_texts: List[str] = []
    total = len(images)
    for i, img in enumerate(images, start=1):
        logging.info("OCR %s page %d/%d", input_file.name, i, total)
        text = ocr_image(img, iam_token, folder_id, api_key=api_key)
        page_texts.append(text or "[no text]")
        document.add_heading(f"Page {i}", level=2)
        document.add_paragraph(page_texts[-1])
        print(f"\r{input_file.name}: {i}/{total} pages processed", end="", flush=True)

    print()
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)
    logging.info("Saved DOCX to %s", output_docx)
    shutil.rmtree(tmp_dir, ignore_errors=True)
    return page_texts


def save_combined_results(all_results: List[Tuple[Path, List[str]]], output_dir: Path) -> None:
    """Save aggregated OCR results to DOCX, TXT and CSV files."""
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    txt_lines: List[str] = []

    for file_path, texts in all_results:
        doc.add_heading(file_path.name, level=1)
        txt_lines.append(file_path.name)
        for page_num, text in enumerate(texts, start=1):
            doc.add_heading(f"Page {page_num}", level=2)
            doc.add_paragraph(text)
            txt_lines.append(f"Page {page_num}\n{text}\n")

    docx_path = output_dir / "all_text.docx"
    doc.save(docx_path)

    txt_path = output_dir / "all_text.txt"
    with open(txt_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(txt_lines))

    csv_path = output_dir / "all_text.csv"
    with open(csv_path, "w", newline="", encoding="utf-8") as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(["file", "page", "text"])
        for file_path, texts in all_results:
            for page_num, text in enumerate(texts, start=1):
                writer.writerow([file_path.name, page_num, text])

    logging.info("Saved combined DOCX to %s", docx_path)
    logging.info("Saved combined TXT to %s", txt_path)
    logging.info("Saved combined CSV to %s", csv_path)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run OCR using Yandex Cloud Vision API")
    parser.add_argument(
        "input_path",
        nargs="?",
        type=Path,
        help="Path to image/PDF or directory with files",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("result"),
        help="Directory to save result DOCX files",
    )
    parser.add_argument(
        "--tmp-dir",
        type=Path,
        default=Path("result") / "tmp",
        help="Temporary directory base",
    )
    parser.add_argument(
        "--iam-token",
        default=os.getenv("YANDEX_IAM_TOKEN"),
        help="Yandex IAM token",
    )
    parser.add_argument(
        "--api-key",
        default=os.getenv("YANDEX_API_KEY"),
        help="Yandex Cloud API Key (alternative to IAM token)",
    )
    parser.add_argument(
        "--folder-id",
        default=os.getenv("YANDEX_FOLDER_ID") or DEFAULT_FOLDER_ID,
        help="Yandex Cloud folder ID",
    )
    parser.add_argument(
        "--merge-output",
        action="store_true",
        help="Save results of all files into single DOCX, TXT and CSV",
    )
    return parser.parse_args()


def collect_gui_args(args: argparse.Namespace) -> argparse.Namespace:
    """Collect missing path arguments using a simple Tk based GUI."""
    import tkinter as tk
    from tkinter import filedialog

    root = tk.Tk()
    root.withdraw()

    if args.input_path is None:
        path_str = filedialog.askdirectory(title="Select directory with files")
        if not path_str:
            path_str = filedialog.askopenfilename(title="Select image or PDF")
        if not path_str:
            raise SystemExit("No input selected")
        args.input_path = Path(path_str)

    if args.output_dir == Path("result"):
        out_dir = filedialog.askdirectory(title="Select output directory")
        if out_dir:
            args.output_dir = Path(out_dir)

    root.destroy()
    return args


def _resolve_credentials(args: argparse.Namespace) -> Tuple[Optional[str], Optional[str]]:
    """Determine which credentials to use: IAM token or API key."""
    iam_token = args.iam_token
    api_key = args.api_key

    # Priority: explicit IAM -> explicit API key -> OAuth->IAM
    if iam_token:
        return iam_token, None
    if api_key:
        return None, api_key

    # Try to obtain IAM via OAuth using helper module
    iam_via_oauth = api.get_iam_token()
    if iam_via_oauth:
        return iam_via_oauth, None

    return None, None


def main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s", stream=sys.stdout)
    args = parse_args()

    if args.input_path is None:
        args = collect_gui_args(args)

    iam_token, api_key = _resolve_credentials(args)
    if not iam_token and not api_key:
        raise SystemExit(
            "No credentials provided. Use one of:\n"
            "  --iam-token <IAM_TOKEN>\n"
            "  --api-key <API_KEY>\n"
            "or set env vars YANDEX_IAM_TOKEN / YANDEX_API_KEY. "
            "Optionally set YANDEX_OAUTH_TOKEN to auto-fetch IAM."
        )

    args.output_dir.mkdir(parents=True, exist_ok=True)

    try:
        files = find_input_files(args.input_path)
    except FileNotFoundError as exc:
        raise SystemExit(str(exc))
    if not files:
        raise SystemExit("No input files found")

    if not args.folder_id:
        raise SystemExit("Folder ID is required. Provide --folder-id or set YANDEX_FOLDER_ID.")

    combined_results: List[Tuple[Path, List[str]]] = []

    for idx, file in enumerate(files, start=1):
        logging.info("Processing file %d/%d", idx, len(files))
        tmp_dir = args.tmp_dir / file.stem
        output_docx = args.output_dir / file.stem / f"{file.stem}.docx"
        try:
            texts = process_file(file, output_docx, tmp_dir, iam_token, args.folder_id, api_key)
            logging.info("Saved OCR result for %s to %s", file, output_docx)
            if args.merge_output:
                combined_results.append((file, texts))
        except Exception:
            logging.exception("Failed to process %s", file)

    if args.merge_output and combined_results:
        save_combined_results(combined_results, args.output_dir)


if __name__ == "__main__":
    main()
